#!/usr/bin/env python3
"""Render the limited Markdown used by DeepSpec experiment reports to DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TABLE_SEPARATOR = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    return parser.parse_args()


def set_east_asia_font(run, name: str) -> None:
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), name
    )


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_text(document: Document, text: str, *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Microsoft YaHei")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if not TABLE_SEPARATOR.match(line)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            text = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.font.size = Pt(8)
            set_east_asia_font(run, "Microsoft YaHei")
            if row_index == 0:
                run.bold = True
                shade_cell(cell, "D9EAF7")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(31, 78, 121)

    code_style = document.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.paragraph_format.left_indent = Cm(0.5)
    code_style.paragraph_format.space_after = Pt(1)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DeepSpec · DeepSeek-V4 DSpark · FSDP2 overlap optimization")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def render(markdown: str) -> Document:
    document = Document()
    configure_document(document)
    document.core_properties.title = "DeepSeek-V4 DSpark FSDP2 overlap optimization"
    document.core_properties.subject = "Single-node 8-GPU overlap tuning and verification"

    lines = markdown.splitlines()
    index = 0
    first_heading = True
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_text(document, "".join(paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_text(document, "\n".join(code_lines), style="Code Block")
        elif stripped.startswith("#"):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if level == 1 and not first_heading:
                document.add_page_break()
            paragraph = document.add_heading(title, level=min(level, 3))
            if first_heading:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].font.size = Pt(22)
                first_heading = False
        elif stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            index -= 1
        elif re.match(r"^-\s+", stripped):
            flush_paragraph()
            add_text(document, re.sub(r"^-\s+", "", stripped), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            add_text(document, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif not stripped:
            flush_paragraph()
        else:
            paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()
    return document


def main() -> None:
    args = parse_args()
    document = render(args.markdown.read_text())
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
